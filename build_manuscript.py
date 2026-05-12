"""
Build a full-length manuscript as a Word document (.docx).
Loads pre-computed statistics from outputs/table1_stats.json and
phase2_experimentation/nodes/stage3/metrics.json, then calls Claude
to generate each section via the claude_call() helper.

Replace all placeholder strings (marked with "replace") with values
specific to your study before running.
"""

import os, sys, json
import numpy as np
import pandas as pd

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from claude_client import claude_call

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORK_DIR  = os.path.dirname(os.path.abspath(__file__))
OUT_PATH  = os.path.join(WORK_DIR, "phase3_writeup", "Manuscript.docx")   # replace filename
NODES_DIR = os.path.join(WORK_DIR, "phase2_experimentation", "nodes")

# ── Load pre-computed stats ───────────────────────────────────────────────────
with open(os.path.join(WORK_DIR, "outputs", "table1_stats.json")) as f:
    stats = json.load(f)

with open(os.path.join(NODES_DIR, "stage3", "metrics.json")) as f:
    cox_metrics = json.load(f)

with open(os.path.join(NODES_DIR, "stage1", "metrics.json")) as f:
    uni_metrics = json.load(f)

# ── Build study context from loaded metrics ───────────────────────────────────
n_total  = cox_metrics.get("n", "N/A")
n_events = cox_metrics.get("events", "N/A")
c_index  = cox_metrics.get("c_index", "N/A")
main_hr  = cox_metrics.get("hr", "N/A")
hr_lo    = cox_metrics.get("hr_lo", "N/A")
hr_hi    = cox_metrics.get("hr_hi", "N/A")
p_val    = cox_metrics.get("p_value", "N/A")

# Replace the text below with a description of YOUR study.
# The section-generation prompts reference CONTEXT, so be specific
# about your study population, primary outcome, and key results.
CONTEXT = f"""
STUDY: <Replace with your study title>

KEY STATISTICS (loaded from metrics files — do not invent numbers):
- Multivariable Cox cohort: N={n_total:,}, events={n_events:,}, C-index={c_index}
- Primary exposure HR={main_hr} [95% CI {hr_lo}–{hr_hi}], p={p_val}

<Add any additional study-specific context here, e.g.:
- Sample size and follow-up period
- Exposure/outcome definitions
- Key covariate results from cox_metrics["subgroups"]
- Extended follow-up results if available
- Clinical narrative and interpretation>
"""


def gen(prompt, max_tokens=2500):
    full = (
        "IMPORTANT: Output plain prose only. No markdown headers. "
        "No bullet points unless specified.\n\n"
        f"{CONTEXT}\n\n{prompt}"
    )
    return claude_call(full, max_tokens=max_tokens)


# ── Generate text sections ────────────────────────────────────────────────────
# Replace each prompt below with instructions appropriate for your study.

print("Generating Introduction...")
intro = gen(
    """
Write a full Introduction section (~600 words, 4 paragraphs).

Para 1: <Replace — background on the clinical problem and disease burden>
Para 2: <Replace — existing evidence and knowledge gap>
Para 3: <Replace — rationale for this study>
Para 4: <Replace — study aim, cohort description, and primary question>
""",
    2000,
)

print("Generating Methods...")
methods = gen(
    """
Write a full Methods section (~900 words). Plain text subheadings:
Study Design and Setting, Study Population, Outcomes, Covariates,
Statistical Analysis, Ethics.

<Replace — fill in your study design, eligibility criteria, outcome
definitions, covariate list, and statistical methods>
""",
    2500,
)

print("Generating Results...")
results = gen(
    """
Write a full Results section (~1,100 words). Plain text subheadings.

Use ALL exact numbers from the context. Reference Table 1 for baseline,
Table 2 for Cox results, Figure 1 and Figure 2 for survival and forest plots.

<Replace — describe your study population, primary unadjusted and
adjusted results, subgroup findings, and any extended follow-up analyses>
""",
    2500,
)

print("Generating Discussion...")
discussion = gen(
    """
Write a full Discussion section (~1,000 words). 5 paragraphs.

Para 1: <Replace — principal finding, contrast with unadjusted analysis>
Para 2: <Replace — comparison with existing literature>
Para 3: <Replace — mechanistic interpretation of key findings>
Para 4: <Replace — subgroup insights>
Para 5: <Replace — limitations and clinical implications>
""",
    2500,
)

print("Generating Conclusion...")
conclusion = gen(
    """
Write a 4-sentence Conclusions paragraph summarising the principal
finding, key secondary findings, clinical implications, and future
research directions.

<Replace — tailor to your study's results>
""",
    400,
)

# ── Build Table 2 (Cox regression) from loaded metrics ───────────────────────
# Replace variable labels and populate HR/CI/p from cox_metrics as needed.
cox_summary = cox_metrics.get("summary", {})  # if stored as a dict

cox_table = [("Variable", "HR", "95% CI", "p-value")]
# Example row — replace with your actual variables:
# cox_table.append(("Primary exposure", str(main_hr), f"{hr_lo}–{hr_hi}", str(p_val)))
# Add remaining covariate rows here...

# ── Subgroup table from metrics ───────────────────────────────────────────────
sg = cox_metrics.get("subgroups", {})

subgroup_table = [("Subgroup", "n", "HR", "95% CI", "p-value")]
# Replace key strings with your subgroup labels as saved in metrics.json
for label, key in [
    # ("Label shown in table", "key_in_subgroups_dict"),
]:
    if key in sg:
        v = sg[key]
        subgroup_table.append((
            label,
            f"{v['n']:,}",
            f"{v['HR']:.2f}",
            f"{v['lo']:.2f}–{v['hi']:.2f}",
            f"{v['p']:.3f}",
        ))

# ── Extended follow-up table ──────────────────────────────────────────────────
ext = cox_metrics.get("extended_followup", {})

extended_table = [("Outcome", "N", "Events", "HR", "95% CI", "p-value")]
# Primary row — replace label:
extended_table.append((
    "Primary outcome",
    f"{n_total:,}", f"{n_events:,}",
    str(main_hr), f"{hr_lo}–{hr_hi}", str(p_val),
))
for label, key in [
    # ("Secondary outcome 1", "key_in_extended_followup"),
]:
    if key in ext:
        v = ext[key]
        extended_table.append((
            label, f"{v['n']:,}", f"{v['events']:,}",
            f"{v['HR']:.3f}", f"{v['lo']:.3f}–{v['hi']:.3f}", f"{v['p']:.3f}",
        ))


# ── Word document ─────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.name = "Times New Roman"
    run.font.bold = True
    run.font.size = Pt(13 if level == 1 else 12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def body(doc, text):
    for para in text.strip().split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if len(para) < 80 and not para.endswith(".") and "\n" not in para:
            p = doc.add_paragraph(para)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
        else:
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.space_after = Pt(6)


def add_table(doc, data, caption, bold_header=True, bold_first_col=False):
    cap = doc.add_paragraph(caption)
    cap.runs[0].bold = True
    cap.runs[0].font.size = Pt(11)
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT

    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            cell.text = str(cell_text)
            run = (cell.paragraphs[0].runs[0]
                   if cell.paragraphs[0].runs
                   else cell.paragraphs[0].add_run(str(cell_text)))
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            if i == 0 and bold_header:
                run.bold = True
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9D9D9")
                tcPr.append(shd)
            if j == 0 and bold_first_col and i > 0:
                run.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()


def add_figure(doc, path, caption, width=Inches(5.5)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Figure not found: {os.path.basename(path)}]")
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].font.size = Pt(10)
        cap.runs[0].italic = True
    doc.add_paragraph()


# ── TITLE PAGE ────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run(
    "<Replace with your manuscript title>"   # replace
)
title_run.bold = True
title_run.font.size = Pt(14)
title_run.font.name = "Times New Roman"

doc.add_paragraph()
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth.add_run(
    "<Author names and affiliations>\n"       # replace
    "Correspondence: <email>\n\n"             # replace
    "Running title: <short title>\n"          # replace
    "Word count: ~5,000\n"
    "Figures: <n>  |  Tables: <n>"            # replace
)
doc.add_page_break()

# ── ABSTRACT ─────────────────────────────────────────────────────────────────
heading(doc, "ABSTRACT")
abstract_text = (
    # Replace with your abstract. Structure: Background / Methods / Results / Conclusions / Keywords.
    "<Replace with your study abstract>\n\n"
    "Keywords: <keyword 1>; <keyword 2>; <keyword 3>"   # replace
)
body(doc, abstract_text)
doc.add_page_break()

# ── BODY SECTIONS ─────────────────────────────────────────────────────────────
heading(doc, "INTRODUCTION")
body(doc, intro)
doc.add_paragraph()

heading(doc, "METHODS")
body(doc, methods)
doc.add_paragraph()

heading(doc, "RESULTS")
body(doc, results)
doc.add_paragraph()

heading(doc, "DISCUSSION")
body(doc, discussion)
doc.add_paragraph()

heading(doc, "CONCLUSIONS", level=2)
body(doc, conclusion)
doc.add_page_break()

# ── TABLES ────────────────────────────────────────────────────────────────────
heading(doc, "TABLES")

# Table 1 — Baseline characteristics
# Replace column headers with your exposure group labels and sample sizes.
t1_rows = [("Characteristic", "Exposed\n(n=?)", "Unexposed\n(n=?)")]   # replace
for r in stats.get("table1_rows", []):
    t1_rows.append((r[0], r[1], r[2]))

add_table(doc, t1_rows,
    "Table 1. Baseline characteristics by exposure status.\n"
    "Data are mean±SD, median [IQR], or n (%).",   # replace caption as needed
    bold_first_col=True)

# Table 2 — Multivariable Cox
add_table(doc, cox_table,
    f"Table 2. Multivariable Cox proportional hazards analysis "
    f"(N={n_total:,}, events={n_events:,}; C-index={c_index}).\n"
    "HR = hazard ratio; CI = confidence interval.",
    bold_first_col=True)

# Table 3 — Subgroup analysis
if len(subgroup_table) > 1:
    add_table(doc, subgroup_table,
        "Table 3. Subgroup analysis — hazard ratio for exposure across "
        "pre-specified clinical subgroups.",
        bold_first_col=True)

# Table 4 — Extended follow-up
if len(extended_table) > 1:
    add_table(doc, extended_table,
        "Table 4. Multivariable Cox analysis at extended follow-up.\n"
        "HR = hazard ratio; CI = confidence interval.",
        bold_first_col=True)

doc.add_page_break()

# ── FIGURES ───────────────────────────────────────────────────────────────────
heading(doc, "FIGURES")

S1 = os.path.join(NODES_DIR, "stage1", "figures")
S3 = os.path.join(NODES_DIR, "stage3", "figures")

# Replace figure paths and captions with your actual figure files.
fig_map = [
    (os.path.join(S1, "km_curve.png"),
     "Figure 1. Kaplan–Meier survival curves by exposure status."),
    (os.path.join(S1, "forest_univariate.png"),
     "Figure 2. Univariate Cox forest plot."),
    (os.path.join(S3, "multivariable_cox.png"),
     "Figure 3. Multivariable Cox results panel."),
    (os.path.join(S3, "subgroup_forest.png"),
     "Figure 4. Subgroup analysis forest plot."),
]

for fpath, caption in fig_map:
    add_figure(doc, fpath, caption)

# ── REFERENCES ────────────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "REFERENCES")

# Replace with references relevant to your study.
refs = [
    "1. <Author(s)>. <Title>. <Journal>. <Year>;<Volume>:<Pages>.",
    "2. <Author(s)>. <Title>. <Journal>. <Year>;<Volume>:<Pages>.",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    p.paragraph_format.space_after = Pt(4)
    if p.runs:
        p.runs[0].font.size = Pt(10)

os.makedirs(os.path.join(WORK_DIR, "phase3_writeup"), exist_ok=True)
doc.save(OUT_PATH)
print(f"\nSaved: {OUT_PATH}")
print(f"File size: {os.path.getsize(OUT_PATH) / 1024:.0f} KB")
